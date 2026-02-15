import os
import json
import subprocess
import secrets
import copy
import threading
import uuid
from functools import wraps
from flask import Flask, request, jsonify, send_file, render_template, session, redirect, url_for
from werkzeug.utils import secure_filename
from openai import OpenAI
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Resolve project paths relative to this file so Flask can find templates/static
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TEMPLATE_DIR = os.path.join(BASE_DIR, 'templates')
STATIC_DIR = os.path.join(BASE_DIR, 'static')

# Load .env from project root so OPENAI_API_KEY is picked up
load_dotenv(os.path.join(BASE_DIR, '.env'))

app = Flask(__name__, template_folder=TEMPLATE_DIR, static_folder=STATIC_DIR)
app.secret_key = os.getenv("SECRET_KEY", secrets.token_hex(32))

# --- PASSWORD PROTECTION ---
APP_PASSWORD = os.getenv("APP_PASSWORD", "").strip()

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if APP_PASSWORD and not session.get('authenticated'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# --- CONFIGURATION ---
UPLOAD_FOLDER = os.path.join(STATIC_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(os.path.dirname(__file__), 'output')
DEFAULT_PHOTO = os.path.join(STATIC_DIR, 'uploads', 'ntrace_logo.jpeg')
ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# --- ASYNC JOB STORAGE ---
jobs: dict[str, dict] = {}

# --- ANONYMIZATION CONSTANTS ---
COMPANY_PHONE = os.getenv("COMPANY_PHONE", "+33 6 62 54 45 33")
COMPANY_EMAIL = "servicecommercial@ntrace-consulting.com"

# --- OPENAI API KEY ---
raw_key = os.getenv("OPENAI_API_KEY") or os.getenv("OPENAI_APIKEY")
if raw_key:
    raw_key = raw_key.strip().strip('"').strip("'")
OPENAI_API_KEY = raw_key
if not OPENAI_API_KEY:
    raise RuntimeError("Set OPENAI_API_KEY in your environment before running the app.")
client = OpenAI(api_key=OPENAI_API_KEY)

# --- LANGUAGE DETECTION ---
FRENCH_KEYWORDS = {
    'le', 'la', 'les', 'des', 'une', 'un', 'de', 'du', 'et', 'ou', 'avec', 'pour', 'sur', 'dans',
    'entreprise', 'compétences', 'expérience', 'formation', 'diplôme', 'poste', 'responsable',
    'gestion', 'développement', 'projet', 'équipe', 'année', 'années', 'mois', 'depuis',
    'janvier', 'février', 'mars', 'avril', 'mai', 'juin', 'juillet', 'août', 'septembre',
    'octobre', 'novembre', 'décembre', 'actuellement', 'présent'
}
ENGLISH_KEYWORDS = {
    'the', 'a', 'an', 'and', 'or', 'with', 'for', 'at', 'in', 'on', 'to', 'of',
    'experience', 'skills', 'education', 'summary', 'degree', 'position', 'manager',
    'development', 'project', 'team', 'year', 'years', 'month', 'months', 'since',
    'january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september',
    'october', 'november', 'december', 'currently', 'present'
}

def detect_language(text: str) -> str:
    words = set(text.lower().split())
    fr_score = len(words & FRENCH_KEYWORDS)
    en_score = len(words & ENGLISH_KEYWORDS)
    return 'en' if en_score > fr_score else 'fr'

def generate_summary(text: str, lang: str) -> str:
    prompt = "Write a concise 2-3 sentence professional profile summary based on this CV. Language: " + ("English." if lang == 'en' else "French.")
    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You write concise CV summaries."},
            {"role": "user", "content": prompt + "\n\nCV Text:\n" + text[:8000]}
        ],
        temperature=0.3,
        max_tokens=180
    )
    return completion.choices[0].message.content.strip()


# ──────────────────────────────────────────────
#  SKILL GROUPING (LLM call)
# ──────────────────────────────────────────────

SKILL_GROUP_PROMPT = """\
You are given a flat list of skills from a CV. Group them into semantic categories \
and return a JSON object where each key is a short category label and the value is \
a comma-separated string of skills belonging to that category.

Rules:
- Use 3-6 categories max. Typical categories: "Compétences techniques", \
"Outils / Logiciels", "Méthodologies", "Compétences fonctionnelles", \
"Langues", "Soft Skills" — but adapt to the actual skills.
- Write category labels in {lang_instruction}.
- Preserve ALL skills — do not drop, rename, or summarize any skill.
- Keep the comma-separated values concise (just the skill names).

Skills:
{skills_json}

Return ONLY valid JSON. Example:
{{"Langages": "Python, Java, C++", "Outils": "Docker, Kubernetes, Git"}}
"""


def group_skills(skills: list, lang: str) -> dict:
    """Use LLM to group a flat skill list into semantic categories."""
    if not skills:
        return {}
    lang_instruction = 'English' if lang == 'en' else 'French'
    prompt = SKILL_GROUP_PROMPT.format(
        lang_instruction=lang_instruction,
        skills_json=json.dumps(skills, ensure_ascii=False),
    )
    try:
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You organize skills into categories. Return only valid JSON."},
                {"role": "user", "content": prompt},
            ],
            temperature=0,
        )
        raw = completion.choices[0].message.content.strip()
        raw = raw.replace("```json", "").replace("```", "").strip()
        grouped = json.loads(raw)
        if isinstance(grouped, dict) and grouped:
            return grouped
    except Exception as e:
        print(f"Skill grouping failed (non-fatal): {e}")
    label = 'Compétences' if lang != 'en' else 'Skills'
    return {label: ', '.join(skills)}


def ensure_schema(data):
    data = data or {}
    pi = data.get('personal_info', {})
    for key in ['name', 'title', 'email', 'phone', 'location', 'summary', 'photo_path']:
        if key == 'photo_path':
            pi.setdefault(key, DEFAULT_PHOTO)
        else:
            pi.setdefault(key, '')
    data['personal_info'] = pi
    data.setdefault('education', [])
    data.setdefault('skills', [])
    data.setdefault('experience', [])
    data.setdefault('language', 'fr')
    return data


# ──────────────────────────────────────────────
#  TEXT EXTRACTION  (PDF / DOCX / DOC)
# ──────────────────────────────────────────────

def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return _extract_pdf(filepath)
    elif ext == '.docx':
        return _extract_docx(filepath)
    elif ext == '.doc':
        return _extract_doc(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def _extract_pdf(filepath):
    reader = PdfReader(filepath)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def _extract_docx(filepath):
    doc = DocxDocument(filepath)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)

def _extract_doc(filepath):
    try:
        result = subprocess.run(
            ['antiword', filepath],
            capture_output=True, text=True, check=True
        )
        return result.stdout
    except FileNotFoundError:
        raise ValueError("Cannot process .doc files: 'antiword' is not installed.")
    except subprocess.CalledProcessError as e:
        raise ValueError(f"Error reading .doc file: {e.stderr}")


# ──────────────────────────────────────────────
#  ANONYMIZATION
# ──────────────────────────────────────────────

def anonymize_data(data):
    anon = copy.deepcopy(data)
    pi = anon.get('personal_info', {})

    # Name → initials (e.g. "Ousmane SY" → "O. S.")
    name = pi.get('name', '').strip()
    if name:
        parts = [p.strip() for p in name.split() if p.strip()]
        if len(parts) >= 2:
            pi['name'] = f"{parts[0][0].upper()}. {parts[-1][0].upper()}."
        elif len(parts) == 1:
            pi['name'] = f"{parts[0][0].upper()}."

    pi['phone'] = COMPANY_PHONE
    pi['email'] = COMPANY_EMAIL
    anon['personal_info'] = pi
    return anon


# ──────────────────────────────────────────────
#  WORD (.docx) DOCUMENT GENERATION
#  Built directly with python-docx for full control.
# ──────────────────────────────────────────────

CLR_BLUE = RGBColor(0x1F, 0x6F, 0xB2)
CLR_GREY = RGBColor(0x64, 0x64, 0x64)
CLR_BLACK = RGBColor(0x00, 0x00, 0x00)
FONT_NAME = 'Trebuchet MS'


def _remove_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _set_table_col_widths(table, col_widths_cm):
    """Force exact column widths by setting tblLayout=fixed, gridCol, AND each cell's tcW.

    col_widths_cm: list of widths in cm, e.g. [2.8, 14.0]
    """
    tbl = table._tbl

    # 1. Set layout to fixed
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # 2. Set total table width
    total_twips = sum(int(w * 567) for w in col_widths_cm)  # 1cm = 567 twips
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.insert(0, tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(total_twips))

    # 3. Set gridCol widths
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gc in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(gc)
    else:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    for w_cm in col_widths_cm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w_cm * 567)))
        tblGrid.append(gc)

    # 4. Set EACH CELL's tcW -- this is what Word actually uses
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            w_twips = int(col_widths_cm[idx] * 567)
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.insert(0, tcW)
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(w_twips))


def _add_run(para, text, size=11, bold=False, italic=False, color=CLR_BLACK, font=FONT_NAME):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def _set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def _add_section_heading(doc, text):
    para = doc.add_paragraph()
    _add_run(para, text, size=12, bold=True, color=CLR_BLUE)
    _set_spacing(para, before=10, after=2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1F6FB2')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _is_real_summary(text: str) -> bool:
    if not text or not text.strip():
        return False
    return len(text.strip()) >= 40


def _add_detail_to_cell(cell, detail):
    p = cell.add_paragraph()
    _add_run(p, '• ', size=10)
    _add_run(p, detail, size=10)
    _set_spacing(p, before=0, after=0, line=13)
    p.paragraph_format.left_indent = Cm(0.3)


def build_cv_document(data, lang):
    """Build a complete CV document using pure python-docx."""
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    pi = data['personal_info']

    # ── HEADER: Logo + Name + Contact ─────────
    header_table = doc.add_table(rows=1, cols=2)
    _remove_borders(header_table)
    _set_table_col_widths(header_table, [3.0, 14.0])

    logo_cell = header_table.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_path = pi.get('photo_path', DEFAULT_PHOTO)
    if logo_path and os.path.exists(logo_path):
        try:
            logo_para.add_run().add_picture(logo_path, width=Cm(2.5))
        except Exception as e:
            print(f"Logo load warning: {e}")

    info_cell = header_table.cell(0, 1)
    name_para = info_cell.paragraphs[0]
    _add_run(name_para, pi.get('name', ''), size=18, bold=True, color=CLR_BLUE)
    _set_spacing(name_para, after=1)

    title_para = info_cell.add_paragraph()
    _add_run(title_para, (pi.get('title') or '').upper(), size=12, bold=True, color=CLR_BLUE)
    _set_spacing(title_para, after=3)

    contact_para = info_cell.add_paragraph()
    phone = pi.get('phone', '')
    email = pi.get('email', '')
    contact_parts = []
    if phone:
        contact_parts.append(f'Tél : {phone}')
    if email:
        contact_parts.append(f'Email : {email}')
    _add_run(contact_para, '\n'.join(contact_parts), size=11)
    _set_spacing(contact_para, after=0)

    # ── Summary ───────────────────────────────
    summary = pi.get('summary', '')
    if _is_real_summary(summary):
        sum_para = doc.add_paragraph()
        _add_run(sum_para, summary, size=11, italic=True, color=CLR_GREY)
        _set_spacing(sum_para, before=6, after=4)

    # ── SKILLS (grouped, comma-separated) ─────
    skills = [s for s in data.get('skills', []) if s and s.strip()]
    if skills:
        skills_title = 'COMPETENCES PROFESSIONNELLES' if lang != 'en' else 'PROFESSIONAL SKILLS'
        _add_section_heading(doc, skills_title)

        grouped = group_skills(skills, lang)
        for category, skill_str in grouped.items():
            p = doc.add_paragraph()
            _add_run(p, f'• {category} : ', size=11, bold=True)
            _add_run(p, skill_str, size=11)
            _set_spacing(p, before=1, after=2, line=14)
            p.paragraph_format.left_indent = Cm(0.3)

    # ── EXPERIENCE (two-column table) ─────────
    experiences = [
        exp for exp in data.get('experience', [])
        if (exp.get('role') and exp['role'].strip()) or (exp.get('company') and exp['company'].strip())
    ]
    if experiences:
        exp_title = 'EXPÉRIENCES PROFESSIONNELLES' if lang != 'en' else 'PROFESSIONAL EXPERIENCE'
        _add_section_heading(doc, exp_title)

        exp_table = doc.add_table(rows=len(experiences), cols=2)
        _remove_borders(exp_table)
        _set_table_col_widths(exp_table, [3.5, 13.3])

        for i, exp in enumerate(experiences):
            period = exp.get('period', '')
            role = exp.get('role', '')
            company = exp.get('company', '')
            details = [d for d in exp.get('details', []) if d and d.strip()]

            left_cell = exp_table.cell(i, 0)
            left_para = left_cell.paragraphs[0]
            _add_run(left_para, period, size=11, bold=True, color=CLR_BLUE)
            _set_spacing(left_para, before=4, after=1)

            if company:
                comp_para = left_cell.add_paragraph()
                _add_run(comp_para, company, size=11)
                _set_spacing(comp_para, before=0, after=2)

            right_cell = exp_table.cell(i, 1)
            role_para = right_cell.paragraphs[0]
            _add_run(role_para, role, size=11, bold=True)
            _set_spacing(role_para, before=4, after=2)

            for detail in details:
                _add_detail_to_cell(right_cell, detail)

    # ── EDUCATION (two-column table) ──────────
    education = [
        edu for edu in data.get('education', [])
        if (edu.get('degree') and edu['degree'].strip()) or (edu.get('school') and edu['school'].strip())
    ]
    if education:
        edu_title = 'FORMATIONS ET DIPLÔMES' if lang != 'en' else 'EDUCATION'
        _add_section_heading(doc, edu_title)

        edu_table = doc.add_table(rows=len(education), cols=2)
        _remove_borders(edu_table)
        _set_table_col_widths(edu_table, [3.5, 13.3])

        for i, edu in enumerate(education):
            period = edu.get('period', '')
            degree = edu.get('degree', '')
            school = edu.get('school', '')
            details = [d for d in edu.get('details', []) if d and d.strip()]

            left_cell = edu_table.cell(i, 0)
            left_para = left_cell.paragraphs[0]
            _add_run(left_para, period, size=11, bold=True, color=CLR_BLUE)
            _set_spacing(left_para, before=4, after=1)

            if school:
                school_para = left_cell.add_paragraph()
                _add_run(school_para, school, size=11)
                _set_spacing(school_para, before=0, after=2)

            right_cell = edu_table.cell(i, 1)
            degree_para = right_cell.paragraphs[0]
            _add_run(degree_para, degree, size=11, bold=True)
            _set_spacing(degree_para, before=4, after=2)

            for detail in details:
                _add_detail_to_cell(right_cell, detail)

    return doc


# ──────────────────────────────────────────────
#  ROUTES
# ──────────────────────────────────────────────

@app.route('/login', methods=['GET', 'POST'])
def login():
    if not APP_PASSWORD:
        return redirect(url_for('index'))
    error = None
    if request.method == 'POST':
        password = request.form.get('password', '')
        if password == APP_PASSWORD:
            session['authenticated'] = True
            return redirect(url_for('index'))
        error = 'Invalid password'
    return render_template('login.html', error=error)

@app.route('/logout')
def logout():
    session.pop('authenticated', None)
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    return render_template('index.html')

EXTRACTION_PROMPT = """\
Extract ALL data from this CV into the JSON structure below. \
Respond in {lang_instruction} and translate content to that language if needed.
If a field is missing, use null.

IMPORTANT RULES:
1. **PRESERVE ALL INFORMATION**: Do NOT summarize, shorten, or omit any details. \
Extract every bullet point, every skill, every responsibility exactly as written.
2. **SKILLS**: Extract all skills mentioned. Keep the original grouping if present.
3. **EXPERIENCE DETAILS**: Include ALL bullet points and responsibilities for each role.
4. **EDUCATION DETAILS**: Include all certifications, coursework, honors, and details.
5. **PERIOD FORMAT**: Use "MM/YYYY - MM/YYYY" or "YYYY - YYYY" as shown in the CV.

Structure:
{{
  "personal_info": {{ "name": "", "title": "", "email": "", "phone": "", "location": "", "summary": "" }},
  "education": [ {{ "period": "YYYY - YYYY", "degree": "", "school": "", "details": [""] }} ],
  "skills": ["skill1", "skill2", "..."],
  "experience": [ {{ "period": "MM/YYYY - MM/YYYY", "role": "", "company": "", "details": ["detail1", "detail2", "..."] }} ]
}}
Return ONLY valid JSON.
"""


def _call_openai_text(raw_text: str, target_lang: str) -> dict:
    lang_instruction = 'English' if target_lang == 'en' else 'French'
    prompt = EXTRACTION_PROMPT.format(lang_instruction=lang_instruction)

    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You convert CV text into structured JSON."},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "text", "text": "CV Text:\n" + raw_text[:100000]},
                ],
            },
        ],
        temperature=0,
    )
    json_str = completion.choices[0].message.content
    json_str = json_str.replace("```json", "").replace("```", "").strip()
    return json.loads(json_str)


# ──────────────────────────────────────────────
#  CV ANALYSIS  (second LLM call)
# ──────────────────────────────────────────────

ANALYSIS_PROMPT = """\
You are an expert CV reviewer. Analyse the following structured CV data and \
return a JSON object with exactly these keys:

1. "candidate_overview" — a concise 2-3 sentence overview of the candidate \
for the reviewer's reference (who they are, their main expertise, \
years/level of experience). This is NOT the same as the CV's "summary" field. \
Write in {lang_instruction}.

2. "missing_fields" — a list of field names that are empty, null, or missing. \
Check these fields: name, title, email, phone, location, summary, education, \
skills, experience. Only list the ones that are actually missing or empty.

3. "suggestions" — an array of objects, one per missing or empty field for \
which you can propose useful content. Each object must have:
  - "field": the field name (e.g. "title", "summary", "location", "skills")
  - "label": a short human-readable label for what this field is. Write in {lang_instruction}.
  - "value": your proposed content for that field.
IMPORTANT: if the CV's "summary" field in personal_info is empty, you MUST \
include a suggestion for it.
Only include suggestions where you can infer reasonable content from the rest \
of the CV. If nothing is missing or you cannot infer a value, return an empty list.

4. "compact_skills" — look at the "skills" array. If the skills are already \
compact (short phrases, keyword-style entries), set this to null. But if any \
skill entry is a long sentence or paragraph (more than ~8 words), rewrite ALL \
the skills as a clean, compact, professional keyword-style list. \
Write in {lang_instruction}. If null, omit the key or set to null.

CV Data:
{cv_json}

Return ONLY valid JSON matching the structure above.
"""


def _call_openai_analysis(cv_data: dict, target_lang: str) -> dict:
    lang_instruction = 'English' if target_lang == 'en' else 'French'
    send_data = copy.deepcopy(cv_data)
    send_data.get('personal_info', {}).pop('photo_path', None)

    prompt = ANALYSIS_PROMPT.format(
        lang_instruction=lang_instruction,
        cv_json=json.dumps(send_data, ensure_ascii=False, indent=2),
    )

    completion = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a professional CV reviewer. Return only valid JSON."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    json_str = completion.choices[0].message.content
    json_str = json_str.replace("```json", "").replace("```", "").strip()
    try:
        result = json.loads(json_str)
    except json.JSONDecodeError:
        result = {"candidate_overview": "", "missing_fields": [], "suggestions": [], "compact_skills": None}

    result.setdefault("candidate_overview", "")
    result.setdefault("missing_fields", [])
    result.setdefault("suggestions", [])
    result.setdefault("compact_skills", None)
    return result


def _process_cv_job(job_id: str, cv_path: str, ext: str):
    try:
        raw_text = extract_text(cv_path)
        target_lang = detect_language(raw_text)
        extracted_data = _call_openai_text(raw_text, target_lang)

        extracted_data['personal_info']['photo_path'] = DEFAULT_PHOTO
        extracted_data['language'] = target_lang

        if not extracted_data['personal_info'].get('summary'):
            extracted_data['personal_info']['summary'] = ''

        try:
            analysis = _call_openai_analysis(extracted_data, target_lang)
        except Exception as e:
            print(f"Analysis call failed (non-fatal): {e}")
            analysis = {
                "summary": "",
                "missing_fields": [],
                "suggestions": [],
                "compact_skills": None,
            }

        extracted_data['analysis'] = analysis
        jobs[job_id] = {"status": "done", "result": extracted_data}

    except Exception as e:
        print(f"CV Parse Error (job {job_id}): {e}")
        jobs[job_id] = {"status": "error", "error": str(e)}


@app.route('/parse-cv', methods=['POST'])
@login_required
def parse_cv():
    if 'cv_file' not in request.files:
        return jsonify({"error": "Missing CV file"}), 400

    cv_file = request.files['cv_file']
    cv_filename = secure_filename(cv_file.filename)
    ext = os.path.splitext(cv_filename)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({
            "error": f"Unsupported format '{ext}'. Please upload a PDF, DOC, or DOCX file."
        }), 400

    cv_path = os.path.join(UPLOAD_FOLDER, cv_filename)
    cv_file.save(cv_path)

    job_id = uuid.uuid4().hex
    jobs[job_id] = {"status": "processing"}

    thread = threading.Thread(
        target=_process_cv_job,
        args=(job_id, cv_path, ext),
        daemon=True,
    )
    thread.start()

    return jsonify({"job_id": job_id}), 202


@app.route('/job-status/<job_id>')
@login_required
def job_status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    if job["status"] == "done":
        result = job["result"]
        del jobs[job_id]
        return jsonify({"status": "done", "result": result})

    if job["status"] == "error":
        err = job["error"]
        del jobs[job_id]
        return jsonify({"status": "error", "error": err}), 500

    return jsonify({"status": "processing"})


@app.route('/generate-docx', methods=['POST'])
@login_required
def generate_docx():
    data = request.json

    try:
        data = ensure_schema(data)
        lang = data.get('language', 'fr')

        data['skills'] = [s for s in data.get('skills', []) if s and s.strip()]
        data['experience'] = [
            exp for exp in data.get('experience', [])
            if (exp.get('role') and exp['role'].strip()) or (exp.get('company') and exp['company'].strip())
        ]
        data['education'] = [
            edu for edu in data.get('education', [])
            if (edu.get('degree') and edu['degree'].strip()) or (edu.get('school') and edu['school'].strip())
        ]
        for exp in data['experience']:
            exp['details'] = [d for d in exp.get('details', []) if d and d.strip()]
        for edu in data['education']:
            edu['details'] = [d for d in edu.get('details', []) if d and d.strip()]

        anon_data = anonymize_data(data)

        # Build Word document directly (no template)
        doc = build_cv_document(anon_data, lang)

        anon_name = anon_data['personal_info'].get('name', 'CV')
        clean_name = ''.join(c for c in anon_name if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')
        filename = f"CV_{clean_name}.docx" if clean_name else "generated_cv.docx"

        docx_path = os.path.join(OUTPUT_FOLDER, "generated_cv.docx")
        doc.save(docx_path)

        return send_file(docx_path, as_attachment=True, download_name=filename)

    except Exception as e:
        print(f"DOCX Gen Error: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    debug_mode = os.getenv("FLASK_DEBUG", "false").lower() == "true"
    app.run(debug=debug_mode, host='0.0.0.0', port=5000)

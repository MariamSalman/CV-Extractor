#!/usr/bin/env python3
"""
build_template.py
-----------------
One-time script that creates  cv_samples/CV_TEMPLATE.docx  from scratch.

The resulting file is a docxtpl-compatible Word document with Jinja2
placeholders.  It visually matches the original TEMPLATE.docx but uses a
simple internal structure (no section breaks, no floating text-boxes).

Run once:   python build_template.py
Then use:   docxtpl renders the template at runtime in app.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# ── Paths ────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
ICONS = os.path.join(BASE, 'static', 'uploads', 'icons')
DST = os.path.join(BASE, 'cv_samples', 'CV_TEMPLATE.docx')

LOGO = os.path.join(ICONS, 'logo.jpeg')
ICON_EMAIL = os.path.join(ICONS, 'icon_email.png')
ICON_PHONE = os.path.join(ICONS, 'icon_phone.png')
ICON_LOC = os.path.join(ICONS, 'icon_location.png')
ICON_DOT = os.path.join(ICONS, 'icon_calendar.png')  # small colored dot

# ── Colours / sizes (from original template measurements) ────────────────
CLR_BLUE = RGBColor(0x00, 0x4A, 0xAC)       # name, headings, period
CLR_TITLE_BLUE = RGBColor(0x05, 0x70, 0xFF)  # title line
CLR_DARK = RGBColor(0x2F, 0x2E, 0x2E)        # degree/role text in period line
CLR_GRAY = RGBColor(0x6F, 0x6F, 0x6F)        # summary, contact text
CLR_BLACK = RGBColor(0x00, 0x00, 0x00)

FONT_NAME = 'Trebuchet MS'

NAME_SIZE = Pt(17)
TITLE_SIZE = Pt(15)
CONTACT_SIZE = Pt(8)
SUMMARY_SIZE = Pt(8)
HEADING_SIZE = Pt(11)
BODY_SIZE = Pt(8)

# Icon sizes (from original)
ICON_EMAIL_W = Emu(143524)
ICON_EMAIL_H = Emu(102519)
ICON_PHONE_W = Emu(91440)
ICON_PHONE_H = Emu(91440)
ICON_LOC_W = Emu(104775)
ICON_LOC_H = Emu(104775)

# Dot icon size (tiny colored bullet before period lines, ~0.04in square)
ICON_DOT_SZ = Emu(33879)

# Logo size (inline version from original)
LOGO_W = Inches(1.66)
LOGO_H = Inches(0.93)

# Page margins (from original: top≈0.31", bot≈0.19", left≈0.39", right≈0.30")
MARGIN_TOP = Cm(0.8)
MARGIN_BOT = Cm(0.5)
MARGIN_LEFT = Cm(1.0)
MARGIN_RIGHT = Cm(0.8)

# Tab stop for period → degree/role alignment
TAB_POS = Emu(2063115)  # 2.26 inches from original

# Indentation values (from original template, in twips → Pt conversion)
INDENT_HEADING = Pt(12)    # ~234 twips – section headings
INDENT_PERIOD = Pt(5)      # period lines (small indent, dot adds visual weight)
INDENT_SCHOOL = Pt(11)     # ~219 twips – school/company names
INDENT_DETAIL = Pt(11)     # detail bullet items base


# ── Helpers ──────────────────────────────────────────────────────────────

def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Set cell margins in twips (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _add_run(para, text, *, font_name=FONT_NAME, size=BODY_SIZE,
             bold=False, color=CLR_BLACK, italic=False):
    """Add a styled run to a paragraph."""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    if italic:
        run.font.italic = True
    # Ensure east-asian / complex-script font is also set
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    return run


def _add_icon(para, icon_path, width, height):
    """Add an inline image (icon) to a paragraph."""
    run = para.add_run()
    run.add_picture(icon_path, width=width, height=height)
    return run


def _add_tab_stop(para, position, alignment='left'):
    """Add a tab stop to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = parse_xml(f'<w:tabs {nsdecls("w")}/>')
        pPr.append(tabs)
    tab_xml = f'<w:tab {nsdecls("w")} w:val="{alignment}" w:pos="{int(position / 635)}"/>'
    tabs.append(parse_xml(tab_xml))


def _set_para_spacing(para, before=None, after=None, line=None):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    if line is not None:
        pf.line_spacing = line


def _set_left_indent(para, indent):
    """Set left indentation on a paragraph."""
    para.paragraph_format.left_indent = indent


def _add_bottom_border(para, color='004AAC', size='12'):
    """Add a colored bottom border to a paragraph (blue line separator)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def _setup_numbering(doc):
    """Create a bullet list numbering definition and return the numId."""
    # Access the numbering part; create if absent
    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part.element

    # Find highest existing abstractNumId and numId
    max_abstract = -1
    for an in numbering_elem.findall(qn('w:abstractNum')):
        aid = int(an.get(qn('w:abstractNumId'), '-1'))
        if aid > max_abstract:
            max_abstract = aid
    max_num = 0
    for n in numbering_elem.findall(qn('w:num')):
        nid = int(n.get(qn('w:numId'), '0'))
        if nid > max_num:
            max_num = nid

    abstract_id = max_abstract + 1
    num_id = max_num + 1

    # Abstract numbering definition – simple bullet (Symbol font dot)
    abstract_xml = (
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_id}">'
        f'  <w:multiLevelType w:val="hybridMultilevel"/>'
        f'  <w:lvl w:ilvl="0">'
        f'    <w:start w:val="1"/>'
        f'    <w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="\u00B7"/>'
        f'    <w:lvlJc w:val="left"/>'
        f'    <w:pPr>'
        f'      <w:ind w:left="720" w:hanging="360"/>'
        f'    </w:pPr>'
        f'    <w:rPr>'
        f'      <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>'
        f'      <w:sz w:val="16"/>'
        f'    </w:rPr>'
        f'  </w:lvl>'
        f'</w:abstractNum>'
    )
    numbering_elem.append(parse_xml(abstract_xml))

    # Num reference
    num_xml = (
        f'<w:num {nsdecls("w")} w:numId="{num_id}">'
        f'  <w:abstractNumId w:val="{abstract_id}"/>'
        f'</w:num>'
    )
    numbering_elem.append(parse_xml(num_xml))

    return num_id


def _add_bullet_para(doc, body, text_or_tag, num_id, *,
                     font_name=FONT_NAME, size=BODY_SIZE, color=CLR_DARK):
    """Add a bulleted paragraph with text (to the document body)."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=Pt(0), after=Pt(0))
    _apply_bullet(para, num_id)
    _add_run(para, text_or_tag, font_name=font_name, size=size, color=color)
    return para


def _apply_bullet(para, num_id):
    """Apply bullet numbering to any paragraph (body or cell)."""
    pPr = para._p.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="0"/>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.append(numPr)


def _add_cell_para(cell, text, *, font_name=FONT_NAME, size=BODY_SIZE,
                   color=CLR_DARK, bold=False, before=Pt(0), after=Pt(0)):
    """Add a styled paragraph to a table cell."""
    para = cell.add_paragraph()
    _set_para_spacing(para, before=before, after=after)
    _add_run(para, text, font_name=font_name, size=size, color=color, bold=bold)
    return para


def _add_hidden_tag(cell_or_doc, tag, *, is_first=False):
    """Add a Jinja2 control tag as a nearly-invisible paragraph.

    If is_first=True, uses the cell's existing first paragraph instead
    of adding a new one.
    """
    if is_first and hasattr(cell_or_doc, 'paragraphs'):
        para = cell_or_doc.paragraphs[0]
    elif hasattr(cell_or_doc, 'add_paragraph'):
        para = cell_or_doc.add_paragraph()
    else:
        para = cell_or_doc.add_paragraph()
    _set_para_spacing(para, before=Pt(0), after=Pt(0), line=Pt(1))
    _add_run(para, tag, size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))
    return para


def _add_jinja_para(doc, tag, *, style=None, font_name=FONT_NAME,
                    size=BODY_SIZE, color=CLR_DARK, bold=False,
                    alignment=None, space_before=None, space_after=None):
    """Add a paragraph with a Jinja2 tag as its text."""
    para = doc.add_paragraph()
    if style:
        para.style = doc.styles[style]
    _add_run(para, tag, font_name=font_name, size=size, color=color, bold=bold)
    if alignment is not None:
        para.alignment = alignment
    _set_para_spacing(para, before=space_before, after=space_after)
    return para


# ── Main build ───────────────────────────────────────────────────────────

def main():
    doc = Document()

    # ── Page setup (A4, narrow margins) ──────────────────────────────────
    section = doc.sections[0]
    section.page_width = Emu(7562850)   # A4
    section.page_height = Emu(10699750)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Remove the default empty paragraph
    if doc.paragraphs:
        doc.element.body.remove(doc.paragraphs[0]._p)

    # ── Enable "Different First Page" header ─────────────────────────────
    section.different_first_page_header_footer = True

    # ── Default header (pages 2+) – logo only ───────────────────────────
    default_header = section.header
    default_header.is_linked_to_previous = False
    logo_only = default_header.paragraphs[0]
    logo_only.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(logo_only, before=Pt(0), after=Pt(0))
    logo_only.add_run().add_picture(LOGO, width=LOGO_W, height=LOGO_H)

    # ── First-page header – logo + name/title/contact side by side ───────
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False

    # Use a borderless 2-column table inside the first-page header
    usable_width = section.page_width - section.left_margin - section.right_margin
    hdr_tbl = first_header.add_table(rows=1, cols=2, width=usable_width)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(hdr_tbl)

    hdr_left = hdr_tbl.cell(0, 0)
    hdr_right = hdr_tbl.cell(0, 1)
    hdr_left.width = Inches(1.9)
    hdr_right.width = Inches(5.5)
    _set_cell_margins(hdr_left, top=0, start=0, bottom=0, end=0)
    _set_cell_margins(hdr_right, top=0, start=60, bottom=0, end=0)

    # Left cell: Logo
    logo_para = hdr_left.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(logo_para, before=Pt(0), after=Pt(0))
    logo_para.add_run().add_picture(LOGO, width=LOGO_W, height=LOGO_H)

    # Right cell: Name
    name_para = hdr_right.paragraphs[0]
    _set_para_spacing(name_para, before=Pt(2), after=Pt(0))
    _add_run(name_para, '{{ name }}', size=NAME_SIZE, bold=True, color=CLR_BLUE)

    # Right cell: Title
    title_para = hdr_right.add_paragraph()
    _set_para_spacing(title_para, before=Pt(0), after=Pt(2))
    _add_run(title_para, '{{ title }}', size=TITLE_SIZE, color=CLR_TITLE_BLUE)

    # Right cell: Contact line – each icon+text conditional via {%r if %}
    contact_para = hdr_right.add_paragraph()
    _set_para_spacing(contact_para, before=Pt(2), after=Pt(0))

    # Email (conditional)
    _add_run(contact_para, '{%r if email %}', size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))
    _add_icon(contact_para, ICON_EMAIL, ICON_EMAIL_W, ICON_EMAIL_H)
    _add_run(contact_para, ' {{ email }}   ', size=CONTACT_SIZE, color=CLR_GRAY)
    _add_run(contact_para, '{%r endif %}', size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))
    # Phone (conditional)
    _add_run(contact_para, '{%r if phone %}', size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))
    _add_icon(contact_para, ICON_PHONE, ICON_PHONE_W, ICON_PHONE_H)
    _add_run(contact_para, ' {{ phone }}   ', size=CONTACT_SIZE, color=CLR_GRAY)
    _add_run(contact_para, '{%r endif %}', size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))
    # Location (conditional)
    _add_run(contact_para, '{%r if location %}', size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))
    _add_icon(contact_para, ICON_LOC, ICON_LOC_W, ICON_LOC_H)
    _add_run(contact_para, ' {{ location }}', size=CONTACT_SIZE, color=CLR_GRAY)
    _add_run(contact_para, '{%r endif %}', size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF))

    # Remove the default empty paragraph the header table may leave above
    for p in first_header.paragraphs:
        if not p.text.strip() and not p.runs:
            p._p.getparent().remove(p._p)

    # ── Summary (conditional – hidden when empty) ──────────────────────
    _add_jinja_para(doc, '{%p if summary %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))
    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_para_spacing(summary_para, before=Pt(6), after=Pt(2), line=1.15)
    _add_run(summary_para, '{{ summary }}', size=SUMMARY_SIZE, color=CLR_GRAY)
    _add_jinja_para(doc, '{%p endif %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))

    # ── Blue separator line ──────────────────────────────────────────────
    sep = doc.add_paragraph()
    _set_para_spacing(sep, before=Pt(0), after=Pt(2))
    _add_bottom_border(sep, color='004AAC', size='12')

    # ── Set up bullet numbering ──────────────────────────────────────────
    num_id = _setup_numbering(doc)

    # ══════════════════════════════════════════════════════════════════════
    #  EDUCATION SECTION  (2-column table: left=period+school, right=degree+details)
    #  Entire section hidden when education list is empty.
    # ══════════════════════════════════════════════════════════════════════
    _add_jinja_para(doc, '{%p if education %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))
    edu_heading = doc.add_paragraph()
    _set_para_spacing(edu_heading, before=Pt(6), after=Pt(4))
    _set_left_indent(edu_heading, INDENT_HEADING)
    _add_run(edu_heading, '{{ education_title }}', size=HEADING_SIZE,
             bold=True, color=CLR_BLUE)
    _add_bottom_border(edu_heading, color='004AAC', size='6')

    # Education table – 3 rows: [loop-start] [content – repeated] [loop-end]
    # docxtpl replaces the entire row containing {%tr ...%}, so the for/endfor
    # must be in their own rows, and the content row in the middle gets looped.
    edu_tbl = doc.add_table(rows=3, cols=2)
    _remove_table_borders(edu_tbl)

    for row in edu_tbl.rows:
        for cell in row.cells:
            cell.width = Inches(2.6) if cell == row.cells[0] else Inches(4.9)
            _set_cell_margins(cell, top=40, start=0, bottom=40, end=20)

    # Row 0 – loop start (entire row is replaced by Jinja2 {% for %})
    _add_hidden_tag(edu_tbl.cell(0, 0),
                    '{%tr for edu in education %}', is_first=True)

    # Row 1 – content template (this row is duplicated per education entry)
    edu_left = edu_tbl.cell(1, 0)
    edu_right = edu_tbl.cell(1, 1)
    _set_cell_margins(edu_left, top=20, start=0, bottom=20, end=20)
    _set_cell_margins(edu_right, top=20, start=20, bottom=20, end=0)

    # -- Left: dot + period, school --
    p_period = edu_left.paragraphs[0]
    _set_para_spacing(p_period, before=Pt(0), after=Pt(0))
    _add_icon(p_period, ICON_DOT, ICON_DOT_SZ, ICON_DOT_SZ)
    _add_run(p_period, ' {{ edu.period }}', size=BODY_SIZE, bold=True, color=CLR_BLUE)
    _add_cell_para(edu_left, '{{ edu.school }}', size=BODY_SIZE, color=CLR_DARK)

    # -- Right: degree, then detail bullets --
    p_degree = edu_right.paragraphs[0]
    _set_para_spacing(p_degree, before=Pt(0), after=Pt(0))
    _add_run(p_degree, '{{ edu.degree }}', size=BODY_SIZE, bold=True, color=CLR_DARK)
    _add_hidden_tag(edu_right, '{%p for detail in edu.details %}')
    p_detail = edu_right.add_paragraph()
    _set_para_spacing(p_detail, before=Pt(0), after=Pt(0))
    _apply_bullet(p_detail, num_id)
    _add_run(p_detail, '{{ detail }}', size=BODY_SIZE, color=CLR_DARK)
    _add_hidden_tag(edu_right, '{%p endfor %}')

    # Row 2 – loop end (entire row is replaced by Jinja2 {% endfor %})
    _add_hidden_tag(edu_tbl.cell(2, 0),
                    '{%tr endfor %}', is_first=True)
    _add_jinja_para(doc, '{%p endif %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))

    # ══════════════════════════════════════════════════════════════════════
    #  SKILLS SECTION  (hidden when skills list is empty)
    # ══════════════════════════════════════════════════════════════════════
    _add_jinja_para(doc, '{%p if skills %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))
    skills_heading = doc.add_paragraph()
    _set_para_spacing(skills_heading, before=Pt(10), after=Pt(4))
    _set_left_indent(skills_heading, INDENT_HEADING)
    _add_run(skills_heading, '{{ skills_title }}', size=HEADING_SIZE,
             bold=True, color=CLR_BLUE)
    _add_bottom_border(skills_heading, color='004AAC', size='6')

    # Skills loop
    _add_jinja_para(doc, '{%p for skill in skills %}',
                    size=Pt(2), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))

    _add_bullet_para(doc, doc.element.body, '{{ skill }}', num_id,
                     size=BODY_SIZE, color=CLR_DARK)

    _add_jinja_para(doc, '{%p endfor %}',
                    size=Pt(2), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))
    _add_jinja_para(doc, '{%p endif %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))

    # ══════════════════════════════════════════════════════════════════════
    #  EXPERIENCE SECTION  (2-column table: left=period+company, right=role+details)
    #  Entire section hidden when experience list is empty.
    # ══════════════════════════════════════════════════════════════════════
    _add_jinja_para(doc, '{%p if experience %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))
    exp_heading = doc.add_paragraph()
    _set_para_spacing(exp_heading, before=Pt(10), after=Pt(4))
    _set_left_indent(exp_heading, INDENT_HEADING)
    _add_run(exp_heading, '{{ experience_title }}', size=HEADING_SIZE,
             bold=True, color=CLR_BLUE)
    _add_bottom_border(exp_heading, color='004AAC', size='6')

    # Experience table – 3 rows: [loop-start] [content – repeated] [loop-end]
    exp_tbl = doc.add_table(rows=3, cols=2)
    _remove_table_borders(exp_tbl)

    for row in exp_tbl.rows:
        for cell in row.cells:
            cell.width = Inches(2.6) if cell == row.cells[0] else Inches(4.9)
            _set_cell_margins(cell, top=40, start=0, bottom=40, end=20)

    # Row 0 – loop start
    _add_hidden_tag(exp_tbl.cell(0, 0),
                    '{%tr for exp in experience %}', is_first=True)

    # Row 1 – content template
    exp_left = exp_tbl.cell(1, 0)
    exp_right = exp_tbl.cell(1, 1)
    _set_cell_margins(exp_left, top=20, start=0, bottom=20, end=20)
    _set_cell_margins(exp_right, top=20, start=20, bottom=20, end=0)

    # -- Left: dot + period, company --
    p_exp_period = exp_left.paragraphs[0]
    _set_para_spacing(p_exp_period, before=Pt(0), after=Pt(0))
    _add_icon(p_exp_period, ICON_DOT, ICON_DOT_SZ, ICON_DOT_SZ)
    _add_run(p_exp_period, ' {{ exp.period }}', size=BODY_SIZE, bold=True, color=CLR_BLUE)
    _add_cell_para(exp_left, '{{ exp.company }}', size=BODY_SIZE, color=CLR_DARK)

    # -- Right: role, then detail bullets --
    p_role = exp_right.paragraphs[0]
    _set_para_spacing(p_role, before=Pt(0), after=Pt(0))
    _add_run(p_role, '{{ exp.role }}', size=BODY_SIZE, bold=True, color=CLR_DARK)
    _add_hidden_tag(exp_right, '{%p for detail in exp.details %}')
    p_exp_detail = exp_right.add_paragraph()
    _set_para_spacing(p_exp_detail, before=Pt(0), after=Pt(0))
    _apply_bullet(p_exp_detail, num_id)
    _add_run(p_exp_detail, '{{ detail }}', size=BODY_SIZE, color=CLR_DARK)
    _add_hidden_tag(exp_right, '{%p endfor %}')

    # Row 2 – loop end
    _add_hidden_tag(exp_tbl.cell(2, 0),
                    '{%tr endfor %}', is_first=True)
    _add_jinja_para(doc, '{%p endif %}',
                    size=Pt(1), color=RGBColor(0xFF, 0xFF, 0xFF),
                    space_before=Pt(0), space_after=Pt(0))

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(DST)
    print(f'✓ Created {DST}')

    # Quick verification
    verify = Document(DST)
    print(f'  Paragraphs: {len(verify.paragraphs)}')
    print(f'  Tables: {len(verify.tables)}')
    print(f'  Sections: {len(verify.sections)}')
    for i, p in enumerate(verify.paragraphs[:5]):
        print(f'  [{i}] {p.text[:60]!r}')


if __name__ == '__main__':
    main()
